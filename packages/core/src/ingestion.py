"""
Document Ingestion System

Multi-format document ingestion with recursive folder scanning.
Supports PDF, DOCX, Markdown, HTML, and plain text.
"""

from dataclasses import dataclass, field
from typing import List, Dict, Optional, Set, Literal
from pathlib import Path
import hashlib
import mimetypes
from datetime import datetime
import re


@dataclass
class Document:
    """Ingested document"""
    doc_id: str
    source_path: str
    content: str
    metadata: Dict[str, any] = field(default_factory=dict)
    content_hash: str = ""
    ingested_at: str = field(default_factory=lambda: datetime.utcnow().isoformat() + "Z")
    
    def __post_init__(self):
        if not self.content_hash:
            self.content_hash = hashlib.sha256(self.content.encode()).hexdigest()


class TextExtractor:
    """Extract text from various file formats"""
    
    @staticmethod
    def extract_txt(file_path: Path) -> str:
        """Extract from plain text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def extract_markdown(file_path: Path) -> str:
        """Extract from Markdown file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Remove markdown formatting (basic cleanup)
        # Remove code blocks
        content = re.sub(r'```[\s\S]*?```', '', content)
        # Remove inline code
        content = re.sub(r'`[^`]+`', '', content)
        # Remove links but keep text
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
        # Remove images
        content = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', '', content)
        # Remove headers markers
        content = re.sub(r'^#+\s+', '', content, flags=re.MULTILINE)
        # Remove bold/italic
        content = re.sub(r'\*\*([^\*]+)\*\*', r'\1', content)
        content = re.sub(r'\*([^\*]+)\*', r'\1', content)
        
        return content.strip()
    
    @staticmethod
    def extract_html(file_path: Path) -> str:
        """Extract from HTML file (basic tag stripping)"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Remove script and style tags
        content = re.sub(r'<script[\s\S]*?</script>', '', content, flags=re.IGNORECASE)
        content = re.sub(r'<style[\s\S]*?</style>', '', content, flags=re.IGNORECASE)
        # Remove HTML tags
        content = re.sub(r'<[^>]+>', '', content)
        # Decode common HTML entities
        content = content.replace('&nbsp;', ' ')
        content = content.replace('&lt;', '<')
        content = content.replace('&gt;', '>')
        content = content.replace('&amp;', '&')
        content = content.replace('&quot;', '"')
        
        return content.strip()
    
    @staticmethod
    def extract_pdf(file_path: Path) -> str:
        """
        Extract from PDF file.
        Requires pypdf or pdfplumber (optional dependency).
        """
        try:
            import pypdf
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                text = []
                for page in reader.pages:
                    text.append(page.extract_text())
                return '\n\n'.join(text)
        except ImportError:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = []
                    for page in pdf.pages:
                        text.append(page.extract_text() or '')
                    return '\n\n'.join(text)
            except ImportError:
                raise ImportError("PDF extraction requires 'pypdf' or 'pdfplumber'. Install with: pip install pypdf")
    
    @staticmethod
    def extract_docx(file_path: Path) -> str:
        """
        Extract from DOCX file.
        Requires python-docx (optional dependency).
        """
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n\n'.join(text)
        except ImportError:
            raise ImportError("DOCX extraction requires 'python-docx'. Install with: pip install python-docx")


class DocumentIngestion:
    """Document ingestion orchestrator"""
    
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.markdown': 'markdown',
        '.html': 'html',
        '.htm': 'html',
        '.pdf': 'pdf',
        '.docx': 'docx',
    }
    
    def __init__(
        self,
        max_file_size_mb: int = 10,
        skip_hidden_files: bool = True,
        skip_binary: bool = True
    ):
        self.max_file_size_mb = max_file_size_mb
        self.skip_hidden_files = skip_hidden_files
        self.skip_binary = skip_binary
        self.extractor = TextExtractor()
    
    def scan_directory(
        self,
        directory: Path,
        recursive: bool = True,
        extensions: Optional[Set[str]] = None
    ) -> List[Path]:
        """
        Scan directory for ingestible files.
        Returns list of file paths.
        """
        directory = Path(directory)
        if not directory.exists():
            raise ValueError(f"Directory does not exist: {directory}")
        
        if not directory.is_dir():
            raise ValueError(f"Not a directory: {directory}")
        
        files = []
        pattern = '**/*' if recursive else '*'
        
        for file_path in directory.glob(pattern):
            if not file_path.is_file():
                continue
            
            # Skip hidden files
            if self.skip_hidden_files and file_path.name.startswith('.'):
                continue
            
            # Check extension
            if extensions:
                if file_path.suffix.lower() not in extensions:
                    continue
            else:
                if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
                    continue
            
            # Check file size
            size_mb = file_path.stat().st_size / (1024 * 1024)
            if size_mb > self.max_file_size_mb:
                continue
            
            files.append(file_path)
        
        return sorted(files)
    
    def ingest_file(self, file_path: Path, metadata: Optional[Dict[str, any]] = None) -> Document:
        """
        Ingest a single file.
        Returns Document with extracted content.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise ValueError(f"File does not exist: {file_path}")
        
        # Check size
        size_mb = file_path.stat().st_size / (1024 * 1024)
        if size_mb > self.max_file_size_mb:
            raise ValueError(f"File exceeds size limit: {size_mb:.2f}MB > {self.max_file_size_mb}MB")
        
        # Detect format
        ext = file_path.suffix.lower()
        file_type = self.SUPPORTED_EXTENSIONS.get(ext, 'unknown')
        
        if file_type == 'unknown':
            raise ValueError(f"Unsupported file type: {ext}")
        
        # Extract content
        try:
            if file_type == 'text':
                content = self.extractor.extract_txt(file_path)
            elif file_type == 'markdown':
                content = self.extractor.extract_markdown(file_path)
            elif file_type == 'html':
                content = self.extractor.extract_html(file_path)
            elif file_type == 'pdf':
                content = self.extractor.extract_pdf(file_path)
            elif file_type == 'docx':
                content = self.extractor.extract_docx(file_path)
            else:
                raise ValueError(f"No extractor for type: {file_type}")
        except Exception as e:
            raise ValueError(f"Extraction failed for {file_path}: {str(e)}")
        
        # Create document
        doc_id = hashlib.sha256(str(file_path).encode()).hexdigest()[:16]
        
        doc_metadata = metadata or {}
        doc_metadata.update({
            'file_type': file_type,
            'file_size_bytes': file_path.stat().st_size,
            'file_name': file_path.name,
            'extension': ext
        })
        
        return Document(
            doc_id=doc_id,
            source_path=str(file_path),
            content=content,
            metadata=doc_metadata
        )
    
    def ingest_directory(
        self,
        directory: Path,
        recursive: bool = True,
        extensions: Optional[Set[str]] = None
    ) -> List[Document]:
        """
        Ingest all files in a directory.
        Returns list of Documents.
        """
        files = self.scan_directory(directory, recursive=recursive, extensions=extensions)
        documents = []
        
        for file_path in files:
            try:
                doc = self.ingest_file(file_path)
                documents.append(doc)
            except Exception as e:
                # Log error but continue
                print(f"Warning: Failed to ingest {file_path}: {e}")
                continue
        
        return documents
    
    def ingest_batch(
        self,
        file_paths: List[Path],
        metadata: Optional[Dict[str, any]] = None
    ) -> List[Document]:
        """
        Ingest multiple specific files.
        Returns list of Documents.
        """
        documents = []
        
        for file_path in file_paths:
            try:
                doc = self.ingest_file(file_path, metadata=metadata)
                documents.append(doc)
            except Exception as e:
                print(f"Warning: Failed to ingest {file_path}: {e}")
                continue
        
        return documents


# Example usage
if __name__ == "__main__":
    from pathlib import Path
    
    ingestion = DocumentIngestion(max_file_size_mb=10)
    
    # Scan directory
    test_dir = Path("./test_docs")
    if test_dir.exists():
        files = ingestion.scan_directory(test_dir, recursive=True)
        print(f"Found {len(files)} files:")
        for f in files[:5]:
            print(f"  - {f}")
        
        # Ingest directory
        docs = ingestion.ingest_directory(test_dir)
        print(f"\nIngested {len(docs)} documents")
        
        for doc in docs[:3]:
            print(f"\nDoc ID: {doc.doc_id}")
            print(f"Source: {doc.source_path}")
            print(f"Type: {doc.metadata.get('file_type')}")
            print(f"Size: {doc.metadata.get('file_size_bytes')} bytes")
            print(f"Content preview: {doc.content[:100]}...")
